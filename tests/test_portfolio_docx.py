# tests/test_portfolio_docx.py

import os
import sqlite3
from pathlib import Path

import pytest
from docx import Document


def _extract_docx_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    return "\n".join(p.text or "" for p in doc.paragraphs)


@pytest.fixture()
def conn():
    # Real connection not used because we monkeypatch data fetchers,
    # but keep signature consistent.
    return sqlite3.connect(":memory:")


def test_export_portfolio_docx_no_projects(monkeypatch, tmp_path, conn):
    """
    When there are no projects, we still produce a DOCX with header + message.
    """
    from src.export import portfolio_docx as mod

    monkeypatch.setattr(mod, "collect_project_data", lambda _conn, _user_id: [])
    monkeypatch.setattr(mod, "get_project_summary_row", lambda *_args, **_kwargs: None)
    monkeypatch.setattr(mod, "get_project_thumbnail_path", lambda *_args, **_kwargs: None)

    docx_path = mod.export_portfolio_to_docx(
        conn=conn,
        user_id=1,
        username="Jordan",
        out_dir=str(tmp_path),
    )

    assert docx_path.exists()
    assert docx_path.suffix.lower() == ".docx"
    assert docx_path.stat().st_size > 0

    text = _extract_docx_text(docx_path)
    # UPDATED: exporter uses hyphen, not em dash
    assert "Portfolio - Jordan" in text
    assert "Generated on" in text
    assert "No projects found. Please analyze some projects first." in text


def test_export_portfolio_docx_happy_path_includes_project_text_and_long_summary(monkeypatch, tmp_path, conn):
    """
    One project: ensures header, project line, and long summary content appear in extracted text.
    """
    from src.export import portfolio_docx as mod

    monkeypatch.setattr(mod, "collect_project_data", lambda _conn, _user_id: [("proj_a", 0.1234)])

    fake_summary = {"display_name": "My Fiction Project"}
    fake_row = {
        "summary": fake_summary,
        "project_type": "code",
        "project_mode": "individual",
        "created_at": "2025-01-01",
    }
    monkeypatch.setattr(mod, "get_project_summary_row", lambda *_args, **_kwargs: fake_row)
    monkeypatch.setattr(mod, "get_project_thumbnail_path", lambda *_args, **_kwargs: None)

    monkeypatch.setattr(mod, "resolve_portfolio_display_name", lambda summary, project_name: "My Fiction Project")

    # UPDATED: exporter now prefers format_date_range() first
    monkeypatch.setattr(mod, "format_date_range", lambda *_args, **_kwargs: "Jan 2025 – Feb 2025")

    # Languages/frameworks are now cleaned via portfolio_helpers functions
    monkeypatch.setattr(mod, "_languages_clean", lambda _summary: "Python, SQL")
    monkeypatch.setattr(mod, "_frameworks_clean", lambda _summary: "FastAPI")

    monkeypatch.setattr(mod, "format_activity_line", lambda *_args, **_kwargs: "Activity: Final 100%")
    monkeypatch.setattr(mod, "strip_percent_tokens", lambda s: (s or "").replace(" 100%", "").strip())

    # Skills: exporter uses _skills_one_line(summary)
    monkeypatch.setattr(mod, "_skills_one_line", lambda _summary: "data analysis, APIs")

    long_tail = " ".join(["verylongtext"] * 50)
    monkeypatch.setattr(mod, "resolve_portfolio_summary_text", lambda _summary: long_tail)
    monkeypatch.setattr(mod, "resolve_portfolio_contribution_bullets", lambda *_args, **_kwargs: ["Did X", "Did Y"])

    docx_path = mod.export_portfolio_to_docx(
        conn=conn,
        user_id=1,
        username="Jordan",
        out_dir=str(tmp_path),
    )

    assert docx_path.exists()
    assert docx_path.stat().st_size > 0

    text = _extract_docx_text(docx_path)

    # UPDATED: exporter uses hyphen, not em dash
    assert "Portfolio - Jordan" in text
    assert "Generated on" in text

    assert "My Fiction Project" in text
    assert "Score:" not in text
    assert "Type:" not in text

    assert "Duration:" in text
    assert "Jan 2025" in text
    assert "Feb 2025" in text

    assert "Languages: Python, SQL" in text
    assert "Frameworks: FastAPI" in text

    assert "Activity: Final" in text

    # Skills (one line now)
    assert "Skills: data analysis, APIs" in text

    assert "Project summary:" in text
    assert "verylongtext" in text

    assert "My contribution:" in text
    assert "Did X" in text
    assert "Did Y" in text


def test_export_portfolio_docx_with_thumbnail_does_not_crash(monkeypatch, tmp_path, conn):
    """
    If a thumbnail path exists, export should still succeed.
    We don't test Word image decoding in depth—only that our code attempts to include it.
    """
    from src.export import portfolio_docx as mod

    monkeypatch.setattr(mod, "collect_project_data", lambda _conn, _user_id: [("proj_a", 0.5)])

    fake_row = {
        "summary": {},
        "project_type": "text",
        "project_mode": "individual",
        "created_at": "2025-01-01",
    }
    monkeypatch.setattr(mod, "get_project_summary_row", lambda *_args, **_kwargs: fake_row)

    # Write a real tiny valid PNG
    thumb_path = tmp_path / "thumb.png"
    thumb_path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        b"\x00\x00\x00\rIHDR"
        b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89"
        b"\x00\x00\x00\x0bIDATx\x9cc``\x00\x00\x00\x02\x00\x01\xe2!\xbc3"
        b"\x00\x00\x00\x00IEND\xaeB`\x82"
    )

    monkeypatch.setattr(mod, "get_project_thumbnail_path", lambda *_args, **_kwargs: str(thumb_path))

    # UPDATED: exporter uses resolvers
    monkeypatch.setattr(mod, "resolve_portfolio_display_name", lambda *_args, **_kwargs: "Project With Thumb")
    monkeypatch.setattr(mod, "format_date_range", lambda *_args, **_kwargs: "N/A")
    monkeypatch.setattr(mod, "format_activity_line", lambda *_args, **_kwargs: "Activity: N/A")
    monkeypatch.setattr(mod, "strip_percent_tokens", lambda s: (s or "").strip())
    monkeypatch.setattr(mod, "_skills_one_line", lambda *_args, **_kwargs: "")
    monkeypatch.setattr(mod, "resolve_portfolio_summary_text", lambda *_args, **_kwargs: "ok")
    monkeypatch.setattr(mod, "resolve_portfolio_contribution_bullets", lambda *_args, **_kwargs: ["ok"])

    docx_path = mod.export_portfolio_to_docx(
        conn=conn,
        user_id=1,
        username="Jordan",
        out_dir=str(tmp_path),
    )

    assert docx_path.exists()
    assert docx_path.stat().st_size > 0

    # Check the docx package has at least one image part
    doc = Document(str(docx_path))
    assert len(doc.part.package.image_parts) >= 1

    text = _extract_docx_text(docx_path)
    assert "Project With Thumb" in text
    assert "Project summary:" in text
    assert "ok" in text


def test_export_portfolio_docx_thumbnail_removed_still_exports(monkeypatch, tmp_path, conn):
    """
    Simulate thumbnail existing, then being removed (path returns None).
    Both exports should succeed.
    """
    from src.export import portfolio_docx as mod

    monkeypatch.setattr(mod, "collect_project_data", lambda _conn, _user_id: [("proj_a", 0.5)])
    fake_row = {
        "summary": {},
        "project_type": "text",
        "project_mode": "individual",
        "created_at": "2025-01-01",
    }
    monkeypatch.setattr(mod, "get_project_summary_row", lambda *_args, **_kwargs: fake_row)

    monkeypatch.setattr(mod, "resolve_portfolio_display_name", lambda *_args, **_kwargs: "Thumb Project")
    monkeypatch.setattr(mod, "format_date_range", lambda *_args, **_kwargs: "N/A")
    monkeypatch.setattr(mod, "format_activity_line", lambda *_args, **_kwargs: "Activity: N/A")
    monkeypatch.setattr(mod, "strip_percent_tokens", lambda s: (s or "").strip())
    monkeypatch.setattr(mod, "_skills_one_line", lambda *_args, **_kwargs: "")
    monkeypatch.setattr(mod, "resolve_portfolio_summary_text", lambda *_args, **_kwargs: "ok")
    monkeypatch.setattr(mod, "resolve_portfolio_contribution_bullets", lambda *_args, **_kwargs: ["ok"])

    # 1) With thumbnail
    thumb_path = tmp_path / "thumb.png"
    thumb_path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        b"\x00\x00\x00\rIHDR"
        b"\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89"
        b"\x00\x00\x00\x0bIDATx\x9cc``\x00\x00\x00\x02\x00\x01\xe2!\xbc3"
        b"\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    monkeypatch.setattr(mod, "get_project_thumbnail_path", lambda *_args, **_kwargs: str(thumb_path))

    docx1 = mod.export_portfolio_to_docx(conn=conn, user_id=1, username="Jordan", out_dir=str(tmp_path))
    assert docx1.exists() and docx1.stat().st_size > 0

    # 2) Thumbnail removed
    monkeypatch.setattr(mod, "get_project_thumbnail_path", lambda *_args, **_kwargs: None)

    docx2 = mod.export_portfolio_to_docx(conn=conn, user_id=1, username="Jordan", out_dir=str(tmp_path))
    assert docx2.exists() and docx2.stat().st_size > 0

    text2 = _extract_docx_text(docx2)
    assert "Thumb Project" in text2
    assert "Project summary:" in text2
    assert "ok" in text2


def test_export_portfolio_docx_reflects_edited_summary_text(monkeypatch, tmp_path, conn):
    """
    If the portfolio summary wording was edited (via overrides),
    the exporter should reflect whatever the CLI formatters output.
    """
    from src.export import portfolio_docx as mod

    monkeypatch.setattr(mod, "collect_project_data", lambda _conn, _user_id: [("proj_a", 0.9)])
    fake_row = {
        "summary": {"portfolio_overrides": {"summary_text": "NEW SUMMARY HERE"}},
        "project_type": "text",
        "project_mode": "individual",
        "created_at": "2025-01-01",
    }
    monkeypatch.setattr(mod, "get_project_summary_row", lambda *_args, **_kwargs: fake_row)
    monkeypatch.setattr(mod, "get_project_thumbnail_path", lambda *_args, **_kwargs: None)

    monkeypatch.setattr(mod, "resolve_portfolio_display_name", lambda *_args, **_kwargs: "Edited Summary Project")
    monkeypatch.setattr(mod, "format_date_range", lambda *_args, **_kwargs: "N/A")
    monkeypatch.setattr(mod, "format_activity_line", lambda *_args, **_kwargs: "Activity: N/A")
    monkeypatch.setattr(mod, "strip_percent_tokens", lambda s: (s or "").strip())
    monkeypatch.setattr(mod, "_skills_one_line", lambda *_args, **_kwargs: "")

    # UPDATED: exporter uses resolver
    monkeypatch.setattr(mod, "resolve_portfolio_summary_text", lambda _summary: "NEW SUMMARY HERE")
    monkeypatch.setattr(mod, "resolve_portfolio_contribution_bullets", lambda *_args, **_kwargs: ["ok"])

    docx_path = mod.export_portfolio_to_docx(conn=conn, user_id=1, username="Jordan", out_dir=str(tmp_path))
    text = _extract_docx_text(docx_path)

    assert "Edited Summary Project" in text
    assert "Project summary:" in text
    assert "NEW SUMMARY HERE" in text


def test_export_portfolio_docx_reflects_edited_contribution_bullets(monkeypatch, tmp_path, conn):
    """
    If contribution bullets were edited, and your CLI summary block includes them,
    ensure the DOCX contains the edited bullets.
    """
    from src.export import portfolio_docx as mod

    monkeypatch.setattr(mod, "collect_project_data", lambda _conn, _user_id: [("proj_a", 0.7)])
    fake_row = {
        "summary": {"portfolio_overrides": {"contribution_bullets": ["Did X", "Did Y"]}},
        "project_type": "code",
        "project_mode": "individual",
        "created_at": "2025-01-01",
    }
    monkeypatch.setattr(mod, "get_project_summary_row", lambda *_args, **_kwargs: fake_row)
    monkeypatch.setattr(mod, "get_project_thumbnail_path", lambda *_args, **_kwargs: None)

    monkeypatch.setattr(mod, "resolve_portfolio_display_name", lambda *_args, **_kwargs: "Edited Bullets Project")
    monkeypatch.setattr(mod, "format_date_range", lambda *_args, **_kwargs: "N/A")
    monkeypatch.setattr(mod, "format_activity_line", lambda *_args, **_kwargs: "Activity: N/A")
    monkeypatch.setattr(mod, "strip_percent_tokens", lambda s: (s or "").strip())
    monkeypatch.setattr(mod, "_skills_one_line", lambda *_args, **_kwargs: "")

    # Keep code metadata deterministic (new exporter functions)
    monkeypatch.setattr(mod, "_languages_clean", lambda _summary: "Python")
    monkeypatch.setattr(mod, "_frameworks_clean", lambda _summary: "None")

    # UPDATED: exporter uses resolvers
    monkeypatch.setattr(mod, "resolve_portfolio_summary_text", lambda *_args, **_kwargs: "Something")
    monkeypatch.setattr(mod, "resolve_portfolio_contribution_bullets", lambda *_args, **_kwargs: ["Did X", "Did Y"])

    docx_path = mod.export_portfolio_to_docx(conn=conn, user_id=1, username="Jordan", out_dir=str(tmp_path))
    text = _extract_docx_text(docx_path)

    assert "Edited Bullets Project" in text
    assert "My contribution:" in text
    assert "Did X" in text
    assert "Did Y" in text


def test_export_portfolio_docx_before_after_edit_display_name_summary_and_contrib_add_then_rewrite(
    monkeypatch, tmp_path, conn
):
    """
    Before/after:
    1) Display name changes (rewrite)
    2) Summary text changes (rewrite)
    3) Contribution bullets: placeholder -> add -> rewrite
    """
    from src.export import portfolio_docx as mod

    # One project
    monkeypatch.setattr(mod, "collect_project_data", lambda _conn, _user_id: [("proj_a", 0.5)])
    monkeypatch.setattr(mod, "get_project_thumbnail_path", lambda *_a, **_k: None)
    monkeypatch.setattr(
        mod,
        "get_project_summary_row",
        lambda *_a, **_k: {
            "summary": {},
            "project_type": "code",
            "project_mode": "individual",
            "created_at": "2025-01-01",
        },
    )

    # Deterministic non-test focus fields
    monkeypatch.setattr(mod, "format_date_range", lambda *_a, **_k: "Jan 2025 – Feb 2025")
    monkeypatch.setattr(mod, "_languages_clean", lambda *_a, **_k: "Python")
    monkeypatch.setattr(mod, "_frameworks_clean", lambda *_a, **_k: "None")
    monkeypatch.setattr(mod, "strip_percent_tokens", lambda s: (s or "").strip())
    monkeypatch.setattr(mod, "format_activity_line", lambda *_a, **_k: "Activity: N/A")
    monkeypatch.setattr(mod, "_skills_one_line", lambda *_a, **_k: "")

    # --- BEFORE ---
    monkeypatch.setattr(mod, "resolve_portfolio_display_name", lambda *_a, **_k: "OLD NAME")
    monkeypatch.setattr(mod, "resolve_portfolio_summary_text", lambda _summary: "OLD SUMMARY")
    monkeypatch.setattr(
        mod,
        "resolve_portfolio_contribution_bullets",
        lambda *_a, **_k: ["[No manual contribution summary provided]"],
    )

    docx1 = mod.export_portfolio_to_docx(conn=conn, user_id=1, username="Jordan", out_dir=str(tmp_path))
    t1 = _extract_docx_text(docx1)

    assert "OLD NAME" in t1
    assert "OLD SUMMARY" in t1
    assert "[No manual contribution summary provided]" in t1

    # --- AFTER (EDIT display name + summary + ADD contribution) ---
    monkeypatch.setattr(mod, "resolve_portfolio_display_name", lambda *_a, **_k: "NEW NAME")
    monkeypatch.setattr(mod, "resolve_portfolio_summary_text", lambda _summary: "NEW SUMMARY")
    monkeypatch.setattr(mod, "resolve_portfolio_contribution_bullets", lambda *_a, **_k: ["Did X"])

    docx2 = mod.export_portfolio_to_docx(conn=conn, user_id=1, username="Jordan", out_dir=str(tmp_path))
    t2 = _extract_docx_text(docx2)

    assert "OLD NAME" not in t2
    assert "NEW NAME" in t2
    assert "OLD SUMMARY" not in t2
    assert "NEW SUMMARY" in t2
    assert "[No manual contribution summary provided]" not in t2
    assert "Did X" in t2

    # --- AFTER (REWRITE contribution) ---
    monkeypatch.setattr(mod, "resolve_portfolio_contribution_bullets", lambda *_a, **_k: ["Did Y"])

    docx3 = mod.export_portfolio_to_docx(conn=conn, user_id=1, username="Jordan", out_dir=str(tmp_path))

    # keep original comments etc everything just make necessary changes
    t3 = _extract_docx_text(docx3)

    assert "Did X" not in t3
    assert "Did Y" in t3

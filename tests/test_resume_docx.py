import json
from pathlib import Path

import pytest
from docx import Document


def _doc_text(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text is not None)


def test_resume_export_happy_json_structure_and_order(monkeypatch, tmp_path):
    """
    Covers: deterministic filename via frozen datetime + structure/order:
    Name -> Contact -> (PROFILE, SKILLS, PROJECTS, EDUCATION & CERTIFICATES)
    Projects are sorted by most-recent end_date/start_date first.
    """
    import src.export.resume_docx as exp
    from docx import Document

    def _doc_text(path: Path) -> str:
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text is not None)

    # Freeze datetime.now() used by exporter
    class _FakeDatetime:
        @staticmethod
        def now():
            class _DT:
                def strftime(self, fmt: str) -> str:
                    if fmt == "%Y-%m-%d_%H-%M-%S":
                        return "2026-01-10_15-36-58"
                    if fmt == "%Y-%m-%d at %H:%M:%S":
                        return "2026-01-10 at 15:36:58"
                    return "2026-01-10_15-36-58"
            return _DT()

    monkeypatch.setattr(exp, "datetime", _FakeDatetime)

    # New-format snapshot (sections + projects w/ role + dates + bullets)
    snapshot = {
        "contact_line": "Phone | Email | LinkedIn | Location",
        "profile_text": "To be updated later.",
        "education_text": "To be updated later.",
        "aggregated_skills": {
            "languages": ["Python 88%", "SQL 12%"], 
            "frameworks": ["FastAPI"],
            "technical_skills": ["Algorithms"],
            "writing_skills": ["Clear communication"],
        },
        "projects": [
            # INTENTIONALLY unsorted input (older first) so we can assert sorting.
            {
                "project_name": "older_project",
                "role": "[Role]",
                "start_date": "2024-12-01",
                "end_date": "2024-12-31",
                "contribution_bullets": ["Old bullet A"],
            },
            {
                "project_name": "newer_project",
                "role": "[Role]",
                "start_date": "2025-08-01",
                "end_date": "2025-11-01",
                "contribution_bullets": ["New bullet A", "New bullet B"],
            },
        ],
    }

    record = {"resume_json": json.dumps(snapshot), "rendered_text": "fallback"}
    out_dir = tmp_path / "out"

    path = exp.export_resume_record_to_docx(username="john", record=record, out_dir=str(out_dir))

    assert out_dir.exists()
    assert path.exists()
    assert path.name == "resume_john_2026-01-10_15-36-58.docx"

    txt = _doc_text(path)

    # Name should appear (exact heading string depends on your implementation;
    # this assertion is flexible and just ensures "john" is present early)
    assert "john" in txt.lower()

    # Required sections exist
    assert "PROFILE" in txt
    assert "SKILLS" in txt
    assert "PROJECTS" in txt
    assert "EDUCATION & CERTIFICATES" in txt

    # Order: PROFILE -> SKILLS -> PROJECTS -> EDUCATION
    idx_profile = txt.find("PROFILE")
    idx_skills = txt.find("SKILLS")
    idx_projects = txt.find("PROJECTS")
    idx_edu = txt.find("EDUCATION & CERTIFICATES")
    assert -1 not in (idx_profile, idx_skills, idx_projects, idx_edu)
    assert idx_profile < idx_skills < idx_projects < idx_edu

    # Skills lines rendered
    assert "Languages:" in txt
    assert ("Python" in txt) and ("SQL" in txt)
    assert "Frameworks:" in txt
    assert "FastAPI" in txt
    assert "Technical skills:" in txt
    assert "Algorithms" in txt
    assert "Writing skills:" in txt
    assert "Clear communication" in txt

    # Projects content: role/date line should show up (format may vary slightly)
    assert "newer_project" in txt
    assert "older_project" in txt
    assert "[Role]" in txt

    # Contributions bullets appear
    assert "New bullet A" in txt
    assert "New bullet B" in txt
    assert "Old bullet A" in txt

    # Sorting check: newer should appear before older
    assert txt.find("newer_project") < txt.find("older_project")

    # Optional: verify date range text appears somewhere (don’t overfit formatting)
    assert "2025" in txt and "2024" in txt


def test_resume_export_nonhappy_no_saved_resumes(monkeypatch, capsys):
    """
    Covers: R2 (flow handler: list_resumes empty)
    """
    import src.menu.resume.flow as flow

    monkeypatch.setattr(flow, "list_resumes", lambda conn, user_id: [])
    ok = flow._handle_export_resume_docx(conn=None, user_id=1, username="john")
    out = capsys.readouterr().out
    assert ok is False
    assert "No saved resumes yet" in out


def test_resume_export_nonhappy_cancel_invalid_selection(monkeypatch, capsys):
    """
    Covers: R3 + R4 (cancel, invalid index)
    """
    import src.menu.resume.flow as flow

    resumes = [{"id": 11, "name": "Resume A", "created_at": "2026-01-01"}]
    monkeypatch.setattr(flow, "list_resumes", lambda conn, user_id: resumes)
    monkeypatch.setattr(
        flow,
        "get_resume_snapshot",
        lambda conn, user_id, rid: {"resume_json": "{}", "rendered_text": ""},
    )
    monkeypatch.setattr(flow, "export_resume_record_to_docx", lambda **k: Path("./out/fake.docx"))

    # cancel (Enter)
    monkeypatch.setattr("builtins.input", lambda _: "")
    ok = flow._handle_export_resume_docx(conn=None, user_id=1, username="john")
    out = capsys.readouterr().out
    assert ok is False
    assert "Cancelled" in out

    # invalid index (999)
    monkeypatch.setattr("builtins.input", lambda _: "999")
    ok = flow._handle_export_resume_docx(conn=None, user_id=1, username="john")
    out = capsys.readouterr().out
    assert ok is False
    assert "Invalid selection" in out


def test_resume_export_nonhappy_record_missing(monkeypatch, capsys):
    """
    Covers: R5 (get_resume_snapshot returns None)
    """
    import src.menu.resume.flow as flow

    resumes = [{"id": 11, "name": "Resume A", "created_at": "2026-01-01"}]
    monkeypatch.setattr(flow, "list_resumes", lambda conn, user_id: resumes)
    monkeypatch.setattr(flow, "get_resume_snapshot", lambda conn, user_id, rid: None)
    monkeypatch.setattr("builtins.input", lambda _: "1")

    ok = flow._handle_export_resume_docx(conn=None, user_id=1, username="john")
    out = capsys.readouterr().out
    assert ok is False
    assert "Unable to load the selected resume" in out


def test_resume_export_fallback_to_rendered_text(monkeypatch, tmp_path):
    """
    Covers: R6 + R7 (malformed JSON -> fallback; missing rendered_text -> message)
    """
    import src.export.resume_docx as exp

    class _FakeDatetime:
        @staticmethod
        def now():
            class _DT:
                def strftime(self, fmt: str) -> str:
                    if fmt == "%Y-%m-%d_%H-%M-%S":
                        return "2026-01-10_15-36-58"
                    if fmt == "%Y-%m-%d at %H:%M:%S":
                        return "2026-01-10 at 15:36:58"
                    return "2026-01-10_15-36-58"
            return _DT()

    monkeypatch.setattr(exp, "datetime", _FakeDatetime)

    out_dir = tmp_path / "out"

    # R6: bad JSON + good rendered_text
    record = {"resume_json": "{not json", "rendered_text": "LINE1\nLINE2\n"}
    path = exp.export_resume_record_to_docx(username="john", record=record, out_dir=str(out_dir))
    txt = _doc_text(path)
    assert "Resume Snapshot" in txt
    assert "LINE1" in txt and "LINE2" in txt

    # R7: bad JSON + missing rendered_text
    record2 = {"resume_json": "{not json", "rendered_text": ""}
    path2 = exp.export_resume_record_to_docx(username="john", record=record2, out_dir=str(out_dir))
    txt2 = _doc_text(path2)
    assert "Resume data is missing or unreadable" in txt2


def test_resume_export_uses_key_role(monkeypatch, tmp_path):
    """Test that export uses resolved key_role instead of [Role] placeholder."""
    import src.export.resume_docx as exp
    from docx import Document

    class _FakeDatetime:
        @staticmethod
        def now():
            class _DT:
                def strftime(self, fmt: str) -> str:
                    return "2026-01-10_15-36-58"
            return _DT()

    monkeypatch.setattr(exp, "datetime", _FakeDatetime)

    snapshot = {
        "aggregated_skills": {},
        "projects": [
            {
                "project_name": "test_project",
                "key_role": "Backend Developer",
                "start_date": "2025-01-01",
                "end_date": "2025-06-01",
                "contribution_bullets": ["Built API endpoints"],
            },
        ],
    }

    record = {"resume_json": json.dumps(snapshot), "rendered_text": ""}
    out_dir = tmp_path / "out"
    path = exp.export_resume_record_to_docx(username="jane", record=record, out_dir=str(out_dir))

    txt = _doc_text(path)
    assert "Backend Developer" in txt
    assert "[Role]" not in txt


def test_resume_export_key_role_override_priority(monkeypatch, tmp_path):
    """Test that resume_key_role_override takes priority over base key_role."""
    import src.export.resume_docx as exp
    from docx import Document

    class _FakeDatetime:
        @staticmethod
        def now():
            class _DT:
                def strftime(self, fmt: str) -> str:
                    return "2026-01-10_15-36-58"
            return _DT()

    monkeypatch.setattr(exp, "datetime", _FakeDatetime)

    snapshot = {
        "aggregated_skills": {},
        "projects": [
            {
                "project_name": "test_project",
                "key_role": "Developer",
                "manual_key_role": "Senior Developer",
                "resume_key_role_override": "Lead Developer",
                "contribution_bullets": ["Led team"],
            },
        ],
    }

    record = {"resume_json": json.dumps(snapshot), "rendered_text": ""}
    out_dir = tmp_path / "out"
    path = exp.export_resume_record_to_docx(username="jane", record=record, out_dir=str(out_dir))

    txt = _doc_text(path)
    # Should use the resume override (highest priority)
    assert "Lead Developer" in txt
    assert "Senior Developer" not in txt
    assert "[Role]" not in txt


def test_resume_export_fallback_to_role_placeholder(monkeypatch, tmp_path):
    """Test that export falls back to [Role] when no key_role is set."""
    import src.export.resume_docx as exp
    from docx import Document

    class _FakeDatetime:
        @staticmethod
        def now():
            class _DT:
                def strftime(self, fmt: str) -> str:
                    return "2026-01-10_15-36-58"
            return _DT()

    monkeypatch.setattr(exp, "datetime", _FakeDatetime)

    snapshot = {
        "aggregated_skills": {},
        "projects": [
            {
                "project_name": "test_project",
                # No key_role set
                "contribution_bullets": ["Did stuff"],
            },
        ],
    }

    record = {"resume_json": json.dumps(snapshot), "rendered_text": ""}
    out_dir = tmp_path / "out"
    path = exp.export_resume_record_to_docx(username="jane", record=record, out_dir=str(out_dir))

    txt = _doc_text(path)
    assert "[Role]" in txt

# src/export/portfolio_docx.py
"""
Export the user's portfolio view to a Word (.docx) document.

Structure:
- NO rank/score/type lines
- Duration formatted as "Mon YYYY – Mon YYYY / Present"
- Languages/frameworks: remove % (optionally threshold languages)
- Activity: remove % token (e.g., "Final 100%" -> "Final")
- Skills: one line
- Project summary + My contribution:
    - labels are OUTER bullets
    - content is NESTED bullets (tight bullet-to-text spacing)
    - no fake leading spaces
- Thumbnails included after project title (left aligned)
- Portfolio edits respected:
  resolve_portfolio_display_name / resolve_portfolio_summary_text /
  resolve_portfolio_contribution_bullets
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re
import sqlite3
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.insights.rank_projects.rank_project_importance import collect_project_data
from src.db import get_project_summary_row, get_project_thumbnail_path
from src.insights.portfolio import (
    format_duration,
    format_activity_line,
    resolve_portfolio_display_name,
    resolve_portfolio_summary_text,
    resolve_portfolio_contribution_bullets,
    get_all_skills_from_summary,
)
from src.services.skill_preferences_service import get_highlighted_skills_for_display

# shared + portfolio helpers
from src.export.shared_helpers import (
    format_date_range,
    strip_percent_tokens,
)
from src.export.portfolio_helpers import reformat_duration_line, _skills_one_line,  _frameworks_clean, _languages_clean
from src.insights.portfolio.formatters import _clean_bullets

# -------------------------
# DOCX helpers (local)
# -------------------------

def _safe_slug(s: str) -> str:
    s = (s or "").strip().lower()
    s = re.sub(r"[^a-z0-9]+", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    return s or "user"


def _add_bullet(doc: Document, text: str) -> None:
    """
    Add an OUTER bullet paragraph in Word.
    """
    p = doc.add_paragraph(text, style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # left-right justify
    p.paragraph_format.space_after = 0


def _add_nested_bullet(doc: Document, text: str) -> None:
    """
    Nested bullet using Word list style.
    """
    p = doc.add_paragraph(text, style="List Bullet 3")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # left-right justify
    p.paragraph_format.space_after = 0

# -------------------------
# Export
# -------------------------

def export_portfolio_to_docx(
    conn: sqlite3.Connection,
    user_id: int,
    username: str,
    out_dir: str = "./out",
) -> Path:
    out_path = Path(out_dir)
    out_path.mkdir(parents=True, exist_ok=True)

    now = datetime.now()
    stamp_filename = now.strftime("%Y-%m-%d_%H-%M-%S")
    stamp_display = now.strftime("%Y-%m-%d at %H:%M:%S")

    filename = f"portfolio_{_safe_slug(username)}_{stamp_filename}.docx"
    filepath = out_path / filename

    project_scores: List[Tuple[str, float]] = collect_project_data(conn, user_id)

    doc = Document()

    # Header
    doc.add_heading(f"Portfolio - {username}", level=0)
    doc.add_paragraph(f"Generated on {stamp_display}")

    if not project_scores:
        doc.add_paragraph("No projects found. Please analyze some projects first.")
        doc.save(str(filepath))
        return filepath

    # Get highlighted skills for portfolio context (applies to all projects)
    highlighted_skills = get_highlighted_skills_for_display(
        conn=conn,
        user_id=user_id,
        context="portfolio",
        context_id=None,
    )

    for project_name, _score in project_scores:
        row = get_project_summary_row(conn, user_id, project_name)
        if row is None:
            continue

        summary: Dict[str, Any] = row.get("summary") or {}
        project_type = row.get("project_type") or summary.get("project_type") or "unknown"
        project_mode = row.get("project_mode") or summary.get("project_mode") or "individual"
        created_at = row.get("created_at") or ""

        # Resolvers (respects portfolio_overrides)
        display_name = resolve_portfolio_display_name(summary, project_name)

        # Project title
        doc.add_heading(str(display_name), level=1)

        # Thumbnail right after title (left aligned)
        thumb = get_project_thumbnail_path(conn, user_id, project_name)
        if thumb:
            p = Path(thumb)
            if p.exists():
                try:
                    doc.add_picture(str(p), width=Inches(2.6))
                except Exception:
                    pass

        # Duration: prefer explicit start/end, fallback to existing formatter output
        date_line = format_date_range(summary.get("start_date"), summary.get("end_date"))
        if not date_line:
            raw_duration = (
                format_duration(project_type, project_mode, created_at, user_id, project_name, conn) or ""
            )
            date_line = reformat_duration_line(raw_duration)
        if not date_line:
            date_line = "N/A"

        # Activity: strip percent tokens and remove prefix
        activity_line = format_activity_line(
            project_type, project_mode, conn, user_id, project_name, summary
        ) or ""
        activity_line = strip_percent_tokens(activity_line)
        if activity_line.lower().startswith("activity:"):
            activity_line = activity_line.split(":", 1)[1].strip()

        # Skills: one line (filtered by skill preferences)
        all_project_skills = get_all_skills_from_summary(summary)
        if highlighted_skills:
            # Filter to only show highlighted skills that exist in this project
            filtered_skills = [s for s in highlighted_skills if s in all_project_skills]
        else:
            # No preferences set - use all skills
            filtered_skills = all_project_skills
        skills_line = ", ".join(filtered_skills) if filtered_skills else ""

        # Summary + contribution (same resolvers as PDF)
        project_summary_text = (resolve_portfolio_summary_text(summary) or "").strip()
        contrib_bullets = resolve_portfolio_contribution_bullets(
            summary,
            project_type,
            project_mode,
            conn,
            user_id,
            project_name,
        ) or []
        contrib_bullets = _clean_bullets(contrib_bullets)

        if not project_summary_text:
            project_summary_text = "[No summary provided]"
        if not contrib_bullets:
            contrib_bullets = ["[No manual contribution summary provided]"]

        # ---------------------------
        # OUTER bullets (consistent)
        # ---------------------------
        _add_bullet(doc, f"Duration: {date_line}")

        if project_type == "code":
            _add_bullet(doc, f"Languages: {_languages_clean(summary)}")
            _add_bullet(doc, f"Frameworks: {_frameworks_clean(summary)}")

        _add_bullet(doc, f"Activity: {activity_line or 'N/A'}")

        if skills_line:
            _add_bullet(doc, f"Skills: {skills_line}")

        # Labels OUTER; content NESTED (proper list geometry)
        _add_bullet(doc, "Project summary:")
        _add_nested_bullet(doc, project_summary_text)

        _add_bullet(doc, "My contribution:")
        for b in contrib_bullets:
            _add_nested_bullet(doc, b)

        # spacing between projects
        doc.add_paragraph("")

    doc.save(str(filepath))
    return filepath
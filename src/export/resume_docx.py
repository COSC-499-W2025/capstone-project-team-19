"""
Export a saved resume snapshot (stored JSON) to a Word (.docx) document.

Output format:
- Name
- Contact
- Profile
- Education
- Skills
- Experience
- Projects
- Certificates

Sections are only shown when they have content.

Saves to ./out/ (created if missing).
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
import json
import re
from typing import Any, Dict, List, Optional

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from src.export.resume_helpers import (
    format_date_range,
    add_section_heading,
    add_bullet,
    add_role_date_line,
    _project_sort_key,
    clean_languages_above_threshold,
    filter_skills_by_highlighted,
)

from src.db import (
    get_contact_parts,
    get_visible_profile_text,
    get_resume_name,
)


def _clean_str(value: Any) -> str | None:
    if not isinstance(value, str):
        return None
    text = value.strip()
    return text or None


def _clean_bullets(values: Any) -> List[str]:
    if not isinstance(values, list):
        return []
    cleaned: List[str] = []
    for item in values:
        text = str(item).strip()
        if not text:
            continue
        if text.startswith(("-", "•")):
            text = text.lstrip("-•").strip()
        if text:
            cleaned.append(text)
    return cleaned


def _resume_display_name(p: Dict[str, Any]) -> str:
    resume_name = _clean_str(p.get("resume_display_name_override"))
    if resume_name:
        return resume_name
    manual_name = _clean_str(p.get("manual_display_name"))
    if manual_name:
        return manual_name
    return p.get("project_name") or "Unnamed project"


def _resume_contribution_bullets(p: Dict[str, Any]) -> List[str]:
    resume_bullets = _clean_bullets(p.get("resume_contributions_override"))
    if resume_bullets:
        return resume_bullets
    manual_bullets = _clean_bullets(p.get("manual_contribution_bullets"))
    if manual_bullets:
        return manual_bullets
    return []


def _resume_key_role(p: Dict[str, Any]) -> str | None:
    resume_role = _clean_str(p.get("resume_key_role_override"))
    if resume_role:
        return resume_role
    manual_role = _clean_str(p.get("manual_key_role"))
    if manual_role:
        return manual_role
    return _clean_str(p.get("key_role"))


def _safe_slug(s: str) -> str:
    s = (s or "").strip().lower()
    s = re.sub(r"[^a-z0-9]+", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    return s or "user"


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 0, line: float | None = None) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line is not None:
        fmt.line_spacing = Pt(line)


def _configure_document_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = Pt(14)

    title_style = doc.styles["Title"]
    title_style.font.name = "Arial"
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(4)
    title_style.paragraph_format.line_spacing = Pt(28)

    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Arial"
    heading1.font.size = Pt(14)
    heading1.font.bold = True
    heading1.paragraph_format.space_before = Pt(0)
    heading1.paragraph_format.space_after = Pt(6)
    heading1.paragraph_format.line_spacing = Pt(18)

    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Arial"
    heading2.font.size = Pt(12)
    heading2.font.bold = True
    heading2.paragraph_format.space_before = Pt(8)
    heading2.paragraph_format.space_after = Pt(2)
    heading2.paragraph_format.line_spacing = Pt(15)

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.25)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(0)
    bullet.paragraph_format.line_spacing = Pt(14)


def _render_education_entries(doc: Document, entries: List[Dict[str, Any]]) -> None:
    for entry in entries:
        title = entry.get("title") or "Untitled"
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(12)
        _set_paragraph_spacing(p, before=8, after=2, line=15)

        details = []
        if entry.get("organization"):
            details.append(str(entry["organization"]).strip())
        if entry.get("date_text"):
            details.append(str(entry["date_text"]).strip())

        if details:
            meta = doc.add_paragraph(" | ".join(details))
            if meta.runs:
                meta.runs[0].italic = True
                meta.runs[0].font.name = "Arial"
                meta.runs[0].font.size = Pt(10.5)
            _set_paragraph_spacing(meta, after=6, line=13)

        description = _clean_str(entry.get("description"))
        if description:
            desc = doc.add_paragraph(description)
            _set_paragraph_spacing(desc, after=2, line=14)

        spacer = doc.add_paragraph("")
        _set_paragraph_spacing(spacer, after=4, line=6)


def _render_experience_entries(doc: Document, entries: List[Dict[str, Any]]) -> None:
    for entry in entries:
        role = entry.get("role") or "Untitled role"
        p = doc.add_paragraph()
        run = p.add_run(role)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(12)
        _set_paragraph_spacing(p, before=8, after=2, line=15)

        details = []
        if entry.get("company"):
            details.append(str(entry["company"]).strip())
        if entry.get("date_text"):
            details.append(str(entry["date_text"]).strip())

        if details:
            meta = doc.add_paragraph(" | ".join(details))
            if meta.runs:
                meta.runs[0].italic = True
                meta.runs[0].font.name = "Arial"
                meta.runs[0].font.size = Pt(10.5)
            _set_paragraph_spacing(meta, after=6, line=13)

        description = _clean_str(entry.get("description"))
        if description:
            desc = doc.add_paragraph(description)
            _set_paragraph_spacing(desc, after=2, line=14)

        spacer = doc.add_paragraph("")
        _set_paragraph_spacing(spacer, after=4, line=6)


def export_resume_record_to_docx(
    *,
    username: str,
    record: Dict[str, Any],
    out_dir: str = "./out",
    highlighted_skills: Optional[List[str]] = None,
    highlighted_skills_by_project: Optional[Dict[str, List[str]]] = None,
    user_profile: Optional[Dict[str, Any]] = None,
    education_entries: Optional[List[Dict[str, Any]]] = None,
    experience_entries: Optional[List[Dict[str, Any]]] = None,
) -> Path:
    out_path = Path(out_dir)
    out_path.mkdir(parents=True, exist_ok=True)

    now = datetime.now()
    stamp_filename = now.strftime("%Y-%m-%d_%H-%M-%S")

    filename = f"resume_{_safe_slug(username)}_{stamp_filename}.docx"
    filepath = out_path / filename

    doc = Document()
    _configure_document_layout(doc)

    profile = user_profile or {}
    contact_parts = get_contact_parts(profile)
    education_entries = education_entries or []
    experience_entries = experience_entries or []

    display_name = get_resume_name(profile, username)
    name_para = doc.add_paragraph(style="Title")
    name_run = name_para.add_run(display_name.upper())
    name_run.font.name = "Arial"
    name_run.font.size = Pt(24)
    name_run.bold = True
    _set_paragraph_spacing(name_para, after=4, line=28)

    contact_bits = []
    if contact_parts["phone"]:
        contact_bits.append(("text", contact_parts["phone"]))
    if contact_parts["email"]:
        contact_bits.append(("text", contact_parts["email"]))
    if contact_parts["linkedin"]:
        contact_bits.append(("link", "LinkedIn", contact_parts["linkedin"]))
    if contact_parts["github"]:
        contact_bits.append(("link", "GitHub", contact_parts["github"]))
    if contact_parts["location"]:
        contact_bits.append(("text", contact_parts["location"]))

    if contact_bits:
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, after=6, line=14)
        for i, item in enumerate(contact_bits):
            if i > 0:
                p.add_run(" | ")

            if item[0] == "text":
                p.add_run(item[1])
            else:
                _, label, url = item
                _add_hyperlink(p, url, label)

    resume_json = record.get("resume_json")
    rendered_text = record.get("rendered_text")

    snapshot: Optional[Dict[str, Any]] = None
    if isinstance(resume_json, str) and resume_json.strip():
        try:
            snapshot = json.loads(resume_json)
        except Exception:
            snapshot = None

    if snapshot is None:
        doc.add_heading("Resume Snapshot", level=1)
        if isinstance(rendered_text, str) and rendered_text.strip():
            for line in rendered_text.splitlines():
                para = doc.add_paragraph(line)
                _set_paragraph_spacing(para, after=2, line=14)
        else:
            para = doc.add_paragraph("Resume data is missing or unreadable.")
            _set_paragraph_spacing(para, after=2, line=14)
        doc.save(str(filepath))
        return filepath

    projects: List[Dict[str, Any]] = snapshot.get("projects") or []
    agg: Dict[str, Any] = snapshot.get("aggregated_skills") or {}

    profile_text = get_visible_profile_text(profile)
    if profile_text:
        add_section_heading(doc, "Profile")
        para = doc.add_paragraph(profile_text)
        _set_paragraph_spacing(para, after=2, line=14)

    education_only = [e for e in education_entries if e.get("entry_type") == "education"]
    certificate_only = [e for e in education_entries if e.get("entry_type") == "certificate"]

    if education_only:
        add_section_heading(doc, "Education")
        _render_education_entries(doc, education_only)

    def add_skill_line(label: str, items: List[str]) -> None:
        clean = [str(x).strip() for x in items if str(x).strip()]
        if not clean:
            return
        para = doc.add_paragraph(f"{label}: {', '.join(sorted(set(clean)))}")
        _set_paragraph_spacing(para, after=2, line=14)

    languages = clean_languages_above_threshold(
        agg.get("languages") or [],
        min_pct=10,
    )

    effective_highlighted = highlighted_skills
    if highlighted_skills_by_project is not None:
        all_hl: set = set()
        for sl in highlighted_skills_by_project.values():
            all_hl.update(sl)
        effective_highlighted = list(all_hl)

    adv = filter_skills_by_highlighted(
        agg.get("advanced") or [],
        effective_highlighted,
    )
    interm = filter_skills_by_highlighted(
        agg.get("intermediate") or [],
        effective_highlighted,
    )
    beg = filter_skills_by_highlighted(
        agg.get("beginner") or [],
        effective_highlighted,
    )
    tech_skills = filter_skills_by_highlighted(
        agg.get("technical_skills") or [],
        effective_highlighted,
    )
    writing_skills = filter_skills_by_highlighted(
        agg.get("writing_skills") or [],
        effective_highlighted,
    )

    add_section_heading(doc, "Skills")
    add_skill_line("Languages", languages)
    add_skill_line("Frameworks", agg.get("frameworks") or [])
    if adv or interm or beg:
        add_skill_line("Advanced", adv)
        add_skill_line("Intermediate", interm)
        add_skill_line("Beginner", beg)
    else:
        add_skill_line("Technical skills", tech_skills)
        add_skill_line("Writing skills", writing_skills)

    if experience_entries:
        add_section_heading(doc, "Experience")
        _render_experience_entries(doc, experience_entries)

    projects_sorted = sorted(
        projects,
        key=_project_sort_key,
        reverse=True,
    )

    add_section_heading(doc, "Projects")

    for p in projects_sorted:
        project_name = _resume_display_name(p)
        heading = doc.add_paragraph(style="Heading 2")
        heading_run = heading.add_run(project_name)
        heading_run.font.name = "Arial"
        heading_run.font.size = Pt(12)
        heading_run.bold = True
        _set_paragraph_spacing(heading, before=8, after=2, line=15)

        role = _resume_key_role(p) or "[Role]"
        date_line = format_date_range(p.get("start_date"), p.get("end_date"))
        add_role_date_line(doc, role, date_line)

        custom_bullets = _resume_contribution_bullets(p)
        contrib_bullets = _clean_bullets(p.get("contribution_bullets") or [])

        bullets_to_use: List[str] = []
        if custom_bullets:
            bullets_to_use = custom_bullets
        elif contrib_bullets:
            bullets_to_use = contrib_bullets
        else:
            if p.get("project_type") == "code":
                activities = p.get("activities") or []
                for act in activities:
                    name = act.get("name") or "activity"
                    top = act.get("top_file")
                    top_info = f" (top: {top})" if top else ""
                    bullets_to_use.append(f"{name}{top_info}")
            else:
                pct = p.get("contribution_percent")
                if isinstance(pct, (int, float)):
                    bullets_to_use.append(f"Contributed to {pct:.1f}% of document")

        for b in bullets_to_use:
            add_bullet(doc, str(b))

        spacer = doc.add_paragraph("")
        _set_paragraph_spacing(spacer, after=4, line=6)

    if certificate_only:
        add_section_heading(doc, "Certificates")
        _render_education_entries(doc, certificate_only)

    doc.save(str(filepath))
    return filepath
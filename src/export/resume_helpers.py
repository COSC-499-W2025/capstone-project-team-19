# src/export/resume_helpers.py
from __future__ import annotations

from datetime import datetime
from typing import Any, Dict, List

from docx import Document

from .shared_helpers import (
    parse_date,
    format_date_range,
    strip_percent_tokens,
    clean_languages_above_threshold,
)

# -------------------------
# Skill filtering helpers
# -------------------------

def filter_skills_by_highlighted(
    skills: List[str],
    highlighted_skills: List[str] | None,
) -> List[str]:
    """Filter display skills by raw highlighted skill names.

    Maps display labels back to raw names for comparison.
    """
    if highlighted_skills is None:
        return skills
    if not skills or not highlighted_skills:
        return []

    tech_skill_map = {
        "architecture_and_design": "Architecture & design",
        "data_structures": "Data structures",
        "frontend_skills": "Frontend development",
        "object_oriented_programming": "Object-oriented programming",
        "security_and_error_handling": "Security & error handling",
        "testing_and_ci": "Testing & CI",
        "algorithms": "Algorithms",
        "backend_development": "Backend development",
        "clean_code_and_quality": "Clean code & quality",
        "devops_and_ci_cd": "DevOps & CI/CD",
        "api_and_backend": "API & backend",
    }
    writing_skill_map = {
        "clarity": "Clear communication",
        "structure": "Structured writing",
        "vocabulary": "Strong vocabulary",
        "argumentation": "Analytical writing",
        "depth": "Critical thinking",
        "process": "Revision & editing",
        "planning": "Planning & organization",
        "research": "Research integration",
        "data_collection": "Data collection",
        "data_analysis": "Data analysis",
    }

    label_to_raw = {}
    for raw, label in tech_skill_map.items():
        label_to_raw[label] = raw
    for raw, label in writing_skill_map.items():
        label_to_raw[label] = raw

    filtered: List[str] = []
    for skill in skills:
        raw_name = label_to_raw.get(skill, skill)
        if raw_name in highlighted_skills:
            filtered.append(skill)
    return filtered

# -------------------------
# Resume-only helpers (DOCX)
# -------------------------

def add_role_date_line(doc: Document, role: str, date_line: str) -> None:
    """
    Renders: Role | Nov 2024 – Dec 2024
    If date_line is empty, still shows role (or placeholder).
    """
    role = (role or "").strip()
    date_line = (date_line or "").strip()

    if role and date_line:
        line = f"{role} | {date_line}"
    elif role:
        line = role
    elif date_line:
        line = date_line
    else:
        line = ""

    if line:
        p = doc.add_paragraph(line)
        if p.runs:
            p.runs[0].italic = True


def _project_sort_key(p: dict) -> datetime:
    """
    Sort by:
      1. end_date (preferred)
      2. start_date
      3. very old fallback (goes last)
    """
    end = parse_date(p.get("end_date"))
    start = parse_date(p.get("start_date"))

    if end:
        return end
    if start:
        return start

    # projects with no dates go last
    return datetime.min


def add_section_heading(doc: Document, title: str) -> None:
    doc.add_heading((title or "").upper(), level=1)


def add_placeholder(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].italic = True


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = 0


__all__ = [
    # shared helpers
    "parse_date",
    "format_date_range",
    "strip_percent_tokens",
    "clean_languages_above_threshold",
    "filter_skills_by_highlighted",
    # resume helpers
    "add_role_date_line",
    "_project_sort_key",
    "add_section_heading",
    "add_placeholder",
    "add_bullet",
]
